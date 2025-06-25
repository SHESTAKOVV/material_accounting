import ast
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell

# ===============================
# Настройки ширины таблицы в процентах
COLUMN_WIDTHS = [9.7, 19.9, 22.8, 17.1, 30.3]

SQL_TYPE_MAPPING = {
    "BigAutoField": ("BIGSERIAL", "auto-increment"),
    "AutoField": ("SERIAL", "auto-increment"),
    "CharField": lambda f: (f"VARCHAR({get_field_arg(f, 'max_length')})", None),
    "TextField": ("TEXT", None),
    "DecimalField": lambda f: (
        f"DECIMAL({get_field_arg(f, 'max_digits')},{get_field_arg(f, 'decimal_places')})", None),
    "IntegerField": ("INTEGER", None),
    "BigIntegerField": ("BIGINT", None),
    "BooleanField": ("BOOLEAN", None),
    "DateField": ("DATE", None),
    "DateTimeField": ("TIMESTAMP", None),
    "FileField": ("FILE", None),
    "ForeignKey": ("BIGINT", None)
}

def get_field_arg(field, arg_name):
    for kw in field.keywords:
        if kw.arg == arg_name:
            if hasattr(kw.value, 'n'):
                return kw.value.n
            elif hasattr(kw.value, 's'):
                return kw.value.s
            elif hasattr(kw.value, 'value'):
                return kw.value.value
    return None

def get_str_field_arg(field, arg_name):
    for kw in field.keywords:
        if kw.arg == arg_name:
            try:
                return ast.literal_eval(kw.value)
            except Exception:
                return None
    return None

def get_field_notes(field, field_type):
    notes = []

    for kw in field.keywords:
        if kw.arg == "auto_now_add" and getattr(kw.value, 'value', False):
            notes.append("auto_now_add=True")
        elif kw.arg == "auto_now" and getattr(kw.value, 'value', False):
            notes.append("auto_now=True")
        elif kw.arg == "default":
            val = get_str_field_arg(field, "default")
            if val is not None:
                notes.append(f"DEFAULT {val}")
        elif kw.arg == "unique" and getattr(kw.value, 'value', False):
            notes.append("UNIQUE")

    if field_type == "ForeignKey":
        for kw in field.keywords:
            if kw.arg == "to":
                related = getattr(kw.value, 'id', None) or getattr(kw.value, 'attr', None)
                if related:
                    notes.append(f"FOREIGN KEY TO {related} ON DELETE CASCADE")
    return ", ".join(notes)

def is_required(field):
    for kw in field.keywords:
        if kw.arg in ("null", "blank") and getattr(kw.value, 'value', False):
            return "N"
    return "Y"

def detect_key(field_type):
    if field_type in ("AutoField", "BigAutoField"):
        return "PK"
    if field_type == "ForeignKey":
        return "FK"
    return ""

def get_sql_type(field, field_type):
    mapped = SQL_TYPE_MAPPING.get(field_type)
    if mapped is None:
        return field_type.upper()
    if callable(mapped):
        return mapped(field)[0]
    return mapped[0]

def parse_models_from_ast(tree):
    models = {}
    for node in tree.body:
        if isinstance(node, ast.ClassDef):
            for base in node.bases:
                if isinstance(base, ast.Attribute) and base.attr == "Model":
                    models[node.name] = node.body
                elif isinstance(base, ast.Name) and base.id == "Model":
                    models[node.name] = node.body
    return models

def add_table_title(doc, index, model_name):
    para = doc.add_paragraph(f"Таблица {index} – Таблица «{model_name}»")
    para.style.font.name = 'PT Astra Serif'
    para.style.font.size = Pt(14)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    return para

def set_cell_text(cell: _Cell, text, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = 'PT Astra Serif'
    run.font.size = Pt(12)
    run.bold = bold
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

def fix_table_layout(table):
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for width in COLUMN_WIDTHS:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width * 50)))  # 1/100 от процента
        tblGrid.append(gridCol)
    tbl.insert(0, tblGrid)

    tblPr = tbl.xpath("./w:tblPr")[0]
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

def create_table_for_model(doc, index, model_name, fields):
    doc.add_paragraph("")  # пустая строка перед
    add_table_title(doc, index, model_name)

    table = doc.add_table(rows=1, cols=5)
    fix_table_layout(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    hdr_cells = table.rows[0].cells
    headers = ["KEY", "FIELD NAME", "DATA TYPE / FIELD SIZE", "REQUIRED?", "NOTES"]
    for i, head in enumerate(headers):
        set_cell_text(hdr_cells[i], head, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)

    for field in fields:
        if not isinstance(field, ast.Assign): continue
        if not isinstance(field.value, ast.Call): continue
        if not isinstance(field.value.func, ast.Attribute): continue

        field_name = field.targets[0].id
        field_type = field.value.func.attr

        key = detect_key(field_type)
        sql_type = get_sql_type(field.value, field_type)
        required = is_required(field.value)
        notes = get_field_notes(field.value, field_type)

        row_cells = table.add_row().cells
        set_cell_text(row_cells[0], key)
        set_cell_text(row_cells[1], field_name)
        set_cell_text(row_cells[2], sql_type)
        set_cell_text(row_cells[3], required)
        set_cell_text(row_cells[4], notes)

    doc.add_paragraph("")  # пустая строка после

def generate_data_dictionary(models_path):
    with open(models_path, "r", encoding="utf-8") as f:
        source = f.read()
    tree = ast.parse(source)
    models = parse_models_from_ast(tree)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Pt(85)  # ~3 см
    section.right_margin = Pt(42.5)  # ~1.5 см

    for idx, (model_name, fields) in enumerate(models.items(), start=1):
        create_table_for_model(doc, idx, model_name, fields)

    doc.save("Database_Tables.docx")
    print("Документ сохранён как Database_Tables.docx")

# ========= Запуск =========
if __name__ == "__main__":
    models_file_path = r"E:\Project_python\material_pojject\main\models.py"  # ← путь обновлён
    if not os.path.exists(models_file_path):
        print("Файл models.py не найден!")
    else:
        generate_data_dictionary(models_file_path)